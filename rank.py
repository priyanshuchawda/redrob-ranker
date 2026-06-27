from __future__ import annotations

import argparse
import json
import sys
from datetime import date
from pathlib import Path
from time import perf_counter

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT / "src"))

from redrob_ranker.io import (  # noqa: E402
    build_product_ranking_output,
    iter_candidates,
    write_audit,
    write_debug_scores,
    write_product_csv,
    write_product_json,
    write_product_outputs,
    write_submission,
)
from redrob_ranker.job_understanding import default_role_requirements, parse_job_description  # noqa: E402
from redrob_ranker.scoring import rank_candidates_streaming, rank_scored_candidates, score_candidates  # noqa: E402
from redrob_ranker.schema import load_candidate_records  # noqa: E402


def main() -> int:
    parser = argparse.ArgumentParser(description="Rank candidates with EvidenceGraph Ranker.")
    parser.add_argument("--job", help="Optional job description text file for product/JD-aware ranking")
    parser.add_argument("--candidates", required=True, help="Path to candidate data: JSONL, JSONL.GZ, JSON, or CSV")
    parser.add_argument("--out", help="Legacy challenge CSV output path")
    parser.add_argument("--output", help="Product JSON output path")
    parser.add_argument("--csv-out", help="Product ranked CSV output path")
    parser.add_argument("--top-n", type=int, default=100, help="Number of rows to emit")
    parser.add_argument(
        "--reference-date",
        type=date.fromisoformat,
        default=date(2026, 6, 17),
        help="Dataset snapshot date for recency signals (default: 2026-06-17)",
    )
    parser.add_argument("--debug-out", help="Optional local CSV with all scored candidates and component scores")
    parser.add_argument("--audit-out", help="Optional local CSV for reviewing the emitted top-N candidates")
    parser.add_argument("--audit", action="store_true", help="Write product audit and evidence artifacts under the output directory")
    args = parser.parse_args()

    started = perf_counter()
    if not args.out and not args.output and not args.csv_out:
        parser.error("provide --out for challenge CSV or --output/--csv-out for product outputs")

    product_mode = bool(args.job or args.output or args.csv_out or args.audit)
    data_quality_report = None
    if product_mode:
        output_dir = _product_output_dir(args)
        quality_path = output_dir / "data_quality_report.json"
        records = load_candidate_records(args.candidates, data_quality_report_path=quality_path)
        candidates = [record.to_scoring_dict() for record in records]
        data_quality_report = json.loads(quality_path.read_text(encoding="utf-8"))
    else:
        candidates = iter_candidates(args.candidates)

    role_requirements = (
        parse_job_description(_read_job_description(args.job))
        if args.job
        else default_role_requirements()
    )
    role_for_scoring = role_requirements if args.job else None
    scored = []
    if product_mode or args.debug_out or args.audit_out:
        scored = score_candidates(candidates, reference_date=args.reference_date, role_requirements=role_for_scoring)
        if not product_mode and len(scored) < args.top_n:
            parser.error(f"requested --top-n {args.top_n}, but only {len(scored)} candidates were loaded")
        emit_top_n = min(args.top_n, len(scored)) if product_mode else args.top_n
        ranked = rank_scored_candidates(scored, top_n=emit_top_n)
    else:
        emit_top_n = args.top_n
        ranked = rank_candidates_streaming(
            candidates,
            reference_date=args.reference_date,
            top_n=emit_top_n,
            role_requirements=role_for_scoring,
        )

    if args.out:
        write_submission(ranked, args.out)
    if args.debug_out:
        write_debug_scores(scored, args.debug_out)
    if args.audit_out:
        write_audit(ranked, {row.candidate_id: row for row in scored}, args.audit_out)

    if product_mode:
        elapsed = perf_counter() - started
        payload = build_product_ranking_output(
            scored,
            ranked,
            role_requirements=role_requirements,
            top_n=emit_top_n,
            runtime_seconds=elapsed,
            job_supplied=bool(args.job),
            data_quality_report=data_quality_report,
        )
        output_dir = _product_output_dir(args)
        write_product_outputs(payload, output_dir, csv_filename=(Path(args.csv_out).name if args.csv_out else "ranked_candidates.csv"))
        if args.output:
            write_product_json(payload, args.output)
        if args.csv_out:
            write_product_csv(payload["rankings"], args.csv_out)
        if args.audit:
            write_audit(ranked, {row.candidate_id: row for row in scored}, output_dir / "top_candidates_audit.csv")

    elapsed = perf_counter() - started
    target = args.out or args.output or args.csv_out
    candidate_count = len(scored) if scored else "streamed"
    print(f"Ranked {candidate_count} candidates and wrote {len(ranked)} rows to {target} in {elapsed:.2f}s")
    return 0


def _product_output_dir(args: argparse.Namespace) -> Path:
    if args.output:
        return Path(args.output).parent
    if args.csv_out:
        return Path(args.csv_out).parent
    return Path("outputs")


def _read_job_description(path: str | Path) -> str:
    job_path = Path(path)
    if job_path.suffix.lower() != ".docx":
        return job_path.read_text(encoding="utf-8")
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Reading .docx job descriptions requires python-docx. "
            "Install requirements.txt or provide a plain .txt job file."
        ) from exc
    doc = Document(str(job_path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


if __name__ == "__main__":
    raise SystemExit(main())

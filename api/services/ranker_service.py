from __future__ import annotations

import csv
import html
import io
import sys
import tempfile
import zipfile
from datetime import date
from pathlib import Path
from time import perf_counter
from typing import Any

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from api.schemas import CompareRequest, EvaluateRequest, RankRequest
from redrob_ranker.comparison import compare_scored_candidates
from redrob_ranker.evaluation import evaluate_payload
from redrob_ranker.io import build_product_ranking_output
from redrob_ranker.job_understanding import default_role_requirements, parse_job_description
from redrob_ranker.scoring import rank_scored_candidates, score_candidates
from redrob_ranker.schema import load_candidate_records
from redrob_ranker.trust_audit import build_trust_audit
from api.services.gemini_service import gemini_service


class RankerService:
    def __init__(self) -> None:
        self._latest_payload: dict[str, Any] | None = None
        self._latest_scored: dict[str, Any] = {}
        self._latest_role_requirements: RoleRequirementMatrix | None = None

    def get_demo_data(self) -> dict:
        return {"job_text": _demo_job_text(), "candidates": _demo_candidates()}

    def rank(self, request: RankRequest) -> dict:
        started = perf_counter()
        candidates = request.candidates or _load_demo_candidates()
        matrix = parse_job_description(request.job_text) if request.job_text else default_role_requirements()
        role_for_scoring = matrix if request.job_text else None
        scored = score_candidates(candidates, reference_date=date(2026, 6, 17), role_requirements=role_for_scoring)
        ranked = rank_scored_candidates(scored, top_n=min(request.top_n, len(scored)))
        payload = build_product_ranking_output(
            scored,
            ranked,
            role_requirements=matrix,
            top_n=min(request.top_n, len(scored)),
            runtime_seconds=perf_counter() - started,
            job_supplied=bool(request.job_text),
        )
        payload = gemini_service.enrich_payload(payload, job_text=request.job_text)
        self._latest_payload = payload
        self._latest_scored = {row.candidate_id: row for row in scored}
        self._latest_role_requirements = matrix
        return payload

    def rank_uploaded_candidates(
        self,
        *,
        candidates_filename: str,
        candidates_content: bytes,
        job_text: str | None,
        job_filename: str | None = None,
        job_content: bytes | None = None,
        top_n: int = 50,
    ) -> dict:
        if top_n < 1:
            raise ValueError("top_n must be at least 1")
        if not candidates_filename:
            raise ValueError("Candidate upload must include a filename")

        supplied_job_text = job_text
        if job_content is not None:
            supplied_job_text = _read_uploaded_job_text(job_filename or "", job_content)

        suffix = _upload_suffix(candidates_filename)
        with tempfile.TemporaryDirectory(prefix="evidencegraph-upload-") as temp_dir:
            candidate_path = Path(temp_dir) / f"candidates{suffix}"
            candidate_path.write_bytes(candidates_content)
            try:
                candidates = [
                    record.to_scoring_dict()
                    for record in load_candidate_records(candidate_path)
                ]
            except (OSError, ValueError) as exc:
                raise ValueError(f"Candidate upload could not be parsed: {exc}") from exc

        if not candidates:
            raise ValueError("Candidate upload did not contain any records")
        payload = self.rank(
            RankRequest(
                job_text=supplied_job_text,
                candidates=candidates,
                top_n=min(top_n, len(candidates)),
            )
        )
        payload["metadata"]["input_pipeline"] = {
            "candidate_file": candidates_filename,
            "candidate_file_type": suffix.lstrip("."),
            "candidate_records_loaded": len(candidates),
            "job_source": job_filename or ("typed job description" if supplied_job_text else "default role matrix"),
            "top_n_requested": top_n,
            "top_n_emitted": len(payload.get("rankings", [])),
            "supported_candidate_inputs": ["csv", "json", "jsonl", "ndjson", "jsonl.gz"],
            "supported_job_inputs": ["typed text", "txt", "md", "docx"],
        }
        return payload

    def list_candidates(self) -> dict:
        payload = self.latest_payload()
        return {"candidates": payload.get("rankings", [])}

    def get_candidate(self, candidate_id: str) -> dict | None:
        for row in self.latest_payload().get("rankings", []):
            if row.get("candidate_id") == candidate_id:
                return row
        return None

    def compare(self, request: CompareRequest) -> dict:
        if (
            request.candidates is None
            and request.job_text is None
            and request.candidate_a_id in self._latest_scored
            and request.candidate_b_id in self._latest_scored
        ):
            comparison = compare_scored_candidates(
                self._latest_scored[request.candidate_a_id],
                self._latest_scored[request.candidate_b_id],
                self._latest_role_requirements or default_role_requirements(),
            )
            return _add_ai_semantic_comparison(comparison, self.latest_payload())

        matrix = parse_job_description(request.job_text) if request.job_text else default_role_requirements()
        candidates = request.candidates or _load_demo_candidates()
        scored = score_candidates(
            candidates,
            reference_date=date(2026, 6, 17),
            role_requirements=matrix if request.job_text else None,
        )
        by_id = {row.candidate_id: row for row in scored}
        try:
            comparison = compare_scored_candidates(by_id[request.candidate_a_id], by_id[request.candidate_b_id], matrix)
            payload = build_product_ranking_output(
                scored,
                rank_scored_candidates(scored, top_n=len(scored)),
                role_requirements=matrix,
                top_n=len(scored),
                runtime_seconds=0,
                job_supplied=bool(request.job_text),
            )
            payload = gemini_service.enrich_payload(payload, job_text=request.job_text)
            return _add_ai_semantic_comparison(comparison, payload)
        except KeyError as exc:
            raise ValueError(f"Candidate id not found: {exc.args[0]}") from exc

    def evaluate(self, request: EvaluateRequest) -> dict:
        payload = request.ranking or self.latest_payload()
        return evaluate_payload(payload, labels=request.labels, top_n=request.top_n)

    def trust_audit(self) -> dict:
        audit = build_trust_audit(self.latest_payload())
        audit["ai_summary"] = gemini_service.trust_audit_ai_summary(audit)
        return audit

    def latest_payload(self) -> dict:
        if self._latest_payload is None:
            return self.rank(RankRequest(job_text=_demo_job_text(), candidates=_load_demo_candidates(), top_n=10))
        return self._latest_payload

    def latest_csv(self) -> str:
        payload = self.latest_payload()
        handle = io.StringIO()
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "rank",
                "candidate_id",
                "final_score",
                "fit_score",
                "proof_score",
                "confidence_score",
                "hireability_score",
                "risk_score",
                "main_reason",
            ],
        )
        writer.writeheader()
        for row in payload.get("rankings", []):
            writer.writerow({field: row.get(field, "") for field in writer.fieldnames})
        return handle.getvalue()

    def latest_submission_csv(self) -> str:
        rows = list(self.latest_payload().get("rankings", []))
        handle = io.StringIO()
        writer = csv.DictWriter(handle, fieldnames=["candidate_id", "rank", "score", "reasoning"])
        writer.writeheader()
        for row in rows:
            writer.writerow(
                {
                    "candidate_id": row.get("candidate_id", ""),
                    "rank": row.get("rank", ""),
                    "score": f"{_submission_score(row, rows):.4f}",
                    "reasoning": row.get("main_reason", ""),
                }
            )
        return handle.getvalue()

    def latest_submission_xlsx(self) -> bytes:
        rows = list(csv.reader(io.StringIO(self.latest_submission_csv())))
        return _xlsx_from_rows(rows)


def _load_demo_candidates() -> list[dict]:
    path = ROOT / "data" / "candidates.jsonl"
    if path.exists():
        return [record.to_scoring_dict() for record in load_candidate_records(path)]
    return _demo_candidates()


def _demo_job_text() -> str:
    path = ROOT / "data" / "job.txt"
    if path.exists():
        return path.read_text(encoding="utf-8")
    return (
        "Senior AI Engineer\n"
        "Must have: Python, FastAPI, vector search, ranking, evaluation.\n"
        "Strong signals: shipped retrieval systems, candidate matching, production APIs.\n"
        "Good to have: Kubernetes, MLOps, recommendation systems.\n"
        "5-9 years experience. Location: Pune or Bangalore, India. Notice under 45 days."
    )


def _demo_candidates() -> list[dict]:
    return [
        {
            "candidate_id": "CAND_DEMO_001",
            "profile": {
                "anonymized_name": "Demo Candidate 1",
                "headline": "Senior AI Engineer for search relevance",
                "summary": "Builds retrieval and ranking systems for product teams.",
                "location": "Bangalore, India",
                "country": "India",
                "years_of_experience": 7.0,
                "current_title": "Senior AI Engineer",
                "current_company": "ProductAI",
                "current_industry": "AI/ML",
            },
            "career_history": [
                {
                    "company": "ProductAI",
                    "title": "Senior AI Engineer",
                    "duration_months": 42,
                    "is_current": True,
                    "industry": "AI/ML",
                    "description": "Owned and shipped production semantic search, two-tower candidate matching, cross encoder reranking, NDCG evaluation, and FastAPI serving.",
                }
            ],
            "skills": [
                {"name": "Python", "proficiency": "expert", "duration_months": 72, "endorsements": 30},
                {"name": "FastAPI", "proficiency": "advanced", "duration_months": 36, "endorsements": 12},
                {"name": "Vector Search", "proficiency": "advanced", "duration_months": 36, "endorsements": 20},
            ],
            "redrob_signals": {
                "last_active_date": "2026-06-05",
                "open_to_work_flag": True,
                "recruiter_response_rate": 0.82,
                "notice_period_days": 30,
                "preferred_work_mode": "hybrid",
                "willing_to_relocate": True,
                "profile_completeness_score": 92,
            },
        },
        {
            "candidate_id": "CAND_DEMO_002",
            "profile": {
                "anonymized_name": "Demo Candidate 2",
                "headline": "ML Platform Engineer",
                "summary": "Strong backend and ML platform ownership.",
                "location": "Pune, India",
                "country": "India",
                "years_of_experience": 6.0,
                "current_title": "ML Platform Engineer",
                "current_company": "SaaSCo",
                "current_industry": "SaaS",
            },
            "career_history": [
                {
                    "company": "SaaSCo",
                    "title": "ML Platform Engineer",
                    "duration_months": 48,
                    "is_current": True,
                    "industry": "SaaS",
                    "description": "Built Python APIs, Kubernetes deployments, model serving, monitoring, and data pipelines for ML teams.",
                }
            ],
            "skills": [
                {"name": "Python", "proficiency": "expert", "duration_months": 60, "endorsements": 22},
                {"name": "Kubernetes", "proficiency": "advanced", "duration_months": 30, "endorsements": 10},
                {"name": "MLOps", "proficiency": "advanced", "duration_months": 32, "endorsements": 13},
            ],
            "redrob_signals": {
                "last_active_date": "2026-05-30",
                "open_to_work_flag": True,
                "recruiter_response_rate": 0.68,
                "notice_period_days": 45,
                "preferred_work_mode": "hybrid",
                "willing_to_relocate": True,
                "profile_completeness_score": 86,
            },
        },
    ]


def _upload_suffix(filename: str) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".jsonl.gz"):
        return ".jsonl.gz"
    suffix = Path(lower_name).suffix
    if suffix in {".jsonl", ".ndjson", ".json", ".csv"}:
        return suffix
    raise ValueError(
        "Unsupported candidate upload type. Use JSONL, JSONL.GZ, JSON, or CSV."
    )


def _read_uploaded_job_text(filename: str, content: bytes) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".docx"):
        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError("DOCX job uploads require python-docx. Install requirements.txt.") from exc
        doc = Document(io.BytesIO(content))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if Path(lower_name).suffix not in {".txt", ".md", ""}:
        raise ValueError("Unsupported job upload type. Use TXT, MD, DOCX, or paste the job description.")
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError(f"Job file {filename} must be UTF-8 text") from exc


def _submission_score(row: dict[str, Any], rows: list[dict[str, Any]]) -> float:
    raw_scores = [float((item.get("components") or {}).get("total") or 0.0) for item in rows]
    top_score = max(raw_scores or [0.0])
    row_score = float((row.get("components") or {}).get("total") or 0.0)
    if top_score <= 0:
        return 0.0
    return max(0.0, min(0.99, (row_score / top_score) * 0.99))


def _xlsx_from_rows(rows: list[list[str]]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _XLSX_CONTENT_TYPES)
        archive.writestr("_rels/.rels", _XLSX_RELS)
        archive.writestr("xl/workbook.xml", _XLSX_WORKBOOK)
        archive.writestr("xl/_rels/workbook.xml.rels", _XLSX_WORKBOOK_RELS)
        archive.writestr("xl/styles.xml", _XLSX_STYLES)
        archive.writestr("xl/worksheets/sheet1.xml", _xlsx_sheet(rows))
    return buffer.getvalue()


def _xlsx_sheet(rows: list[list[str]]) -> str:
    row_xml: list[str] = []
    for row_index, row in enumerate(rows, start=1):
        style = ' s="1"' if row_index == 1 else ""
        cells = []
        for col_index, value in enumerate(row, start=1):
            ref = f"{_column_name(col_index)}{row_index}"
            if row_index > 1 and col_index in {2, 3}:
                cells.append(f'<c r="{ref}" s="2"><v>{html.escape(str(value))}</v></c>')
            else:
                cells.append(f'<c r="{ref}" t="inlineStr"{style}><is><t>{html.escape(str(value))}</t></is></c>')
        row_xml.append(f'<row r="{row_index}">{"".join(cells)}</row>')
    dimension = f"A1:D{max(len(rows), 1)}"
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<dimension ref="{dimension}"/>'
        '<sheetViews><sheetView workbookViewId="0"><pane ySplit="1" topLeftCell="A2" activePane="bottomLeft" state="frozen"/></sheetView></sheetViews>'
        '<cols><col min="1" max="1" width="18" customWidth="1"/><col min="2" max="2" width="10" customWidth="1"/><col min="3" max="3" width="12" customWidth="1"/><col min="4" max="4" width="110" customWidth="1"/></cols>'
        f'<sheetData>{"".join(row_xml)}</sheetData>'
        f'<autoFilter ref="{dimension}"/>'
        '</worksheet>'
    )


def _column_name(index: int) -> str:
    name = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name


_XLSX_CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
  <Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>
</Types>"""

_XLSX_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>"""

_XLSX_WORKBOOK = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets><sheet name="Submission" sheetId="1" r:id="rId1"/></sheets>
</workbook>"""

_XLSX_WORKBOOK_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>"""

_XLSX_STYLES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <fonts count="2"><font><sz val="11"/><name val="Calibri"/></font><font><b/><sz val="11"/><color rgb="FFFFFFFF"/><name val="Calibri"/></font></fonts>
  <fills count="3"><fill><patternFill patternType="none"/></fill><fill><patternFill patternType="gray125"/></fill><fill><patternFill patternType="solid"><fgColor rgb="FF10233F"/><bgColor indexed="64"/></patternFill></fill></fills>
  <borders count="1"><border><left/><right/><top/><bottom/><diagonal/></border></borders>
  <cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>
  <cellXfs count="3"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/><xf numFmtId="0" fontId="1" fillId="2" borderId="0" xfId="0" applyFont="1" applyFill="1" applyAlignment="1"><alignment horizontal="center"/></xf><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0" applyAlignment="1"><alignment horizontal="right"/></xf></cellXfs>
  <cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>
</styleSheet>"""


def _add_ai_semantic_comparison(comparison: dict, payload: dict) -> dict:
    rows = {row.get("candidate_id"): row for row in payload.get("rankings", [])}
    a = rows.get((comparison.get("candidate_a") or {}).get("candidate_id"), {})
    b = rows.get((comparison.get("candidate_b") or {}).get("candidate_id"), {})
    comparison["ai_semantic_comparison"] = {
        **gemini_service.signal_fusion_summary(payload),
        "summary": "Gemini assisted insight compares contextual strengths separately from deterministic ranking.",
        "hidden_strengths_difference": list((a.get("ai_contextual_fit") or {}).get("hidden_strengths") or [])
        + list((b.get("ai_contextual_fit") or {}).get("hidden_strengths") or []),
        "risk_difference": [
            *(risk.get("explanation", "") for risk in a.get("risks", []) if isinstance(risk, dict)),
            *(risk.get("explanation", "") for risk in b.get("risks", []) if isinstance(risk, dict)),
        ],
        "interview_checks": list(comparison.get("what_to_verify") or []),
    }
    return comparison

